from PyPDF2 import PdfReader, PdfWriter
from pikepdf import Pdf
from PIL import Image
import io
import zipfile
from typing import List, BinaryIO
import logging

logger = logging.getLogger(__name__)


class PDFService:
    """Service for PDF manipulation operations"""
    
    @staticmethod
    async def merge_pdfs(pdf_files: List[bytes]) -> bytes:
        """
        Merge multiple PDF files into one
        
        Args:
            pdf_files: List of PDF file contents as bytes
            
        Returns:
            Merged PDF as bytes
        """
        try:
            writer = PdfWriter()
            
            for pdf_content in pdf_files:
                reader = PdfReader(io.BytesIO(pdf_content))
                for page in reader.pages:
                    writer.add_page(page)
            
            # Write to bytes
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Merged {len(pdf_files)} PDFs successfully")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error merging PDFs: {e}")
            raise
    
    @staticmethod
    async def split_pdf(pdf_content: bytes, page_ranges: List[int]) -> bytes:
        """
        Extract specific pages from PDF
        
        Args:
            pdf_content: PDF file content as bytes
            page_ranges: List of page indices (0-indexed) to extract
            
        Returns:
            New PDF with selected pages as bytes
        """
        try:
            reader = PdfReader(io.BytesIO(pdf_content))
            writer = PdfWriter()
            
            for page_num in page_ranges:
                if 0 <= page_num < len(reader.pages):
                    writer.add_page(reader.pages[page_num])
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Split PDF with {len(page_ranges)} pages")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error splitting PDF: {e}")
            raise
    
    @staticmethod
    async def compress_pdf(pdf_content: bytes, quality: str = "medium") -> bytes:
        """
        Compress PDF file using pikepdf
        
        Args:
            pdf_content: PDF file content as bytes
            quality: Compression quality (low, medium, high)
            
        Returns:
            Compressed PDF as bytes
        """
        try:
            from pikepdf import ObjectStreamMode
            
            # Quality settings using proper enum values
            quality_settings = {
                "low": {"compress_streams": True, "preserve_pdfa": False, "object_stream_mode": ObjectStreamMode.generate},
                "medium": {"compress_streams": True, "preserve_pdfa": True, "object_stream_mode": ObjectStreamMode.preserve},
                "high": {"compress_streams": True, "preserve_pdfa": True, "object_stream_mode": ObjectStreamMode.disable}
            }
            
            settings = quality_settings.get(quality, quality_settings["medium"])
            
            # Open and compress
            pdf = Pdf.open(io.BytesIO(pdf_content))
            output = io.BytesIO()
            
            pdf.save(
                output,
                compress_streams=settings["compress_streams"],
                preserve_pdfa=settings["preserve_pdfa"],
                object_stream_mode=settings["object_stream_mode"]
            )
            
            output.seek(0)
            compressed_data = output.getvalue()
            
            original_size = len(pdf_content)
            compressed_size = len(compressed_data)
            reduction = ((original_size - compressed_size) / original_size) * 100
            
            logger.info(f"Compressed PDF: {reduction:.1f}% reduction")
            return compressed_data
            
        except Exception as e:
            logger.error(f"Error compressing PDF: {e}")
            raise
    
    @staticmethod
    async def rotate_pdf(pdf_content: bytes, rotation: int, pages: List[int] = None) -> bytes:
        """
        Rotate PDF pages
        
        Args:
            pdf_content: PDF file content as bytes
            rotation: Rotation angle (90, 180, 270)
            pages: List of page indices to rotate (None = all pages)
            
        Returns:
            Rotated PDF as bytes
        """
        try:
            reader = PdfReader(io.BytesIO(pdf_content))
            writer = PdfWriter()
            
            for i, page in enumerate(reader.pages):
                if pages is None or i in pages:
                    page.rotate(rotation)
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Rotated PDF by {rotation} degrees")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error rotating PDF: {e}")
            raise
    
    @staticmethod
    async def reorder_pdf(pdf_content: bytes, page_order: List[int]) -> bytes:
        """
        Reorder PDF pages
        
        Args:
            pdf_content: PDF file content as bytes
            page_order: New order of pages (0-indexed)
            
        Returns:
            Reordered PDF as bytes
        """
        try:
            reader = PdfReader(io.BytesIO(pdf_content))
            writer = PdfWriter()
            
            for page_num in page_order:
                if 0 <= page_num < len(reader.pages):
                    writer.add_page(reader.pages[page_num])
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Reordered PDF with {len(page_order)} pages")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error reordering PDF: {e}")
            raise
    
    @staticmethod
    def get_pdf_info(pdf_content: bytes) -> dict:
        """
        Get PDF metadata
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Dictionary with PDF information
        """
        try:
            reader = PdfReader(io.BytesIO(pdf_content))
            
            return {
                "pages": len(reader.pages),
                "size_bytes": len(pdf_content),
                "metadata": reader.metadata if reader.metadata else {}
            }
            
        except Exception as e:
            logger.error(f"Error getting PDF info: {e}")
            raise

    @staticmethod
    async def pdf_to_word(pdf_content: bytes) -> bytes:
        """
        Convert PDF to Word document (DOCX)
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Word document as bytes
        """
        try:
            from pdf2docx import Converter
            import tempfile
            import os
            
            # Create temporary files for conversion
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
                pdf_temp.write(pdf_content)
                pdf_path = pdf_temp.name
            
            docx_path = pdf_path.replace('.pdf', '.docx')
            
            try:
                # Convert PDF to DOCX
                cv = Converter(pdf_path)
                cv.convert(docx_path)
                cv.close()
                
                # Read the output file
                with open(docx_path, 'rb') as f:
                    docx_content = f.read()
                
                logger.info("Converted PDF to Word successfully")
                return docx_content
                
            finally:
                # Cleanup temp files
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                    
        except Exception as e:
            logger.error(f"Error converting PDF to Word: {e}")
            raise

    @staticmethod
    async def pdf_to_images(pdf_content: bytes, image_format: str = "jpeg", dpi: int = 200) -> bytes:
        """
        Convert PDF pages to images (returns ZIP file with images)
        
        Args:
            pdf_content: PDF file content as bytes
            image_format: Output format (jpeg, png)
            dpi: Resolution in dots per inch
            
        Returns:
            ZIP file containing images as bytes
        """
        try:
            from pdf2image import convert_from_bytes
            
            # Convert PDF to images
            images = convert_from_bytes(pdf_content, dpi=dpi)
            
            # Create ZIP file with images
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for i, image in enumerate(images):
                    img_buffer = io.BytesIO()
                    if image_format.lower() == "png":
                        image.save(img_buffer, format='PNG')
                        ext = 'png'
                    else:
                        image.save(img_buffer, format='JPEG', quality=95)
                        ext = 'jpg'
                    img_buffer.seek(0)
                    zip_file.writestr(f'page_{i + 1}.{ext}', img_buffer.getvalue())
            
            zip_buffer.seek(0)
            logger.info(f"Converted PDF to {len(images)} images")
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            raise

    @staticmethod
    async def images_to_pdf(image_contents: List[bytes]) -> bytes:
        """
        Convert images to PDF
        
        Args:
            image_contents: List of image file contents as bytes
            
        Returns:
            PDF file as bytes
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.pdfgen import canvas
            
            output = io.BytesIO()
            c = canvas.Canvas(output, pagesize=A4)
            a4_width, a4_height = A4
            
            for img_content in image_contents:
                # Open image with PIL
                img = Image.open(io.BytesIO(img_content))
                
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                
                # Get image dimensions
                img_width, img_height = img.size
                
                # Calculate scaling to fit A4 page while maintaining aspect ratio
                width_ratio = a4_width / img_width
                height_ratio = a4_height / img_height
                scale = min(width_ratio, height_ratio) * 0.95  # 5% margin
                
                new_width = img_width * scale
                new_height = img_height * scale
                
                # Center the image on the page
                x = (a4_width - new_width) / 2
                y = (a4_height - new_height) / 2
                
                # Save image to temporary buffer
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='JPEG', quality=95)
                img_buffer.seek(0)
                
                # Create ImageReader from buffer
                from reportlab.lib.utils import ImageReader
                img_reader = ImageReader(img_buffer)
                
                # Draw image on canvas
                c.drawImage(img_reader, x, y, new_width, new_height)
                c.showPage()
            
            c.save()
            output.seek(0)
            
            logger.info(f"Converted {len(image_contents)} images to PDF")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting images to PDF: {e}")
            raise

    @staticmethod
    async def add_watermark(pdf_content: bytes, watermark_text: str, opacity: float = 0.3) -> bytes:
        """
        Add text watermark to PDF
        
        Args:
            pdf_content: PDF file content as bytes
            watermark_text: Text to use as watermark
            opacity: Watermark opacity (0.0 to 1.0)
            
        Returns:
            Watermarked PDF as bytes
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import Color
            
            reader = PdfReader(io.BytesIO(pdf_content))
            writer = PdfWriter()
            
            for page in reader.pages:
                # Get page dimensions
                page_width = float(page.mediabox.width)
                page_height = float(page.mediabox.height)
                
                # Create watermark
                watermark_buffer = io.BytesIO()
                c = canvas.Canvas(watermark_buffer, pagesize=(page_width, page_height))
                
                # Set watermark properties
                c.setFillColor(Color(0.5, 0.5, 0.5, alpha=opacity))
                c.setFont("Helvetica-Bold", 50)
                
                # Rotate and position watermark
                c.saveState()
                c.translate(page_width / 2, page_height / 2)
                c.rotate(45)
                c.drawCentredString(0, 0, watermark_text)
                c.restoreState()
                
                c.save()
                watermark_buffer.seek(0)
                
                # Merge watermark with page
                watermark_reader = PdfReader(watermark_buffer)
                page.merge_page(watermark_reader.pages[0])
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Added watermark '{watermark_text}' to PDF")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error adding watermark: {e}")
            raise

    @staticmethod
    async def add_page_numbers(pdf_content: bytes, position: str = "bottom-center") -> bytes:
        """
        Add page numbers to PDF
        
        Args:
            pdf_content: PDF file content as bytes
            position: Position of page numbers (bottom-center, bottom-right, bottom-left)
            
        Returns:
            PDF with page numbers as bytes
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import black
            
            reader = PdfReader(io.BytesIO(pdf_content))
            writer = PdfWriter()
            total_pages = len(reader.pages)
            
            for i, page in enumerate(reader.pages):
                page_width = float(page.mediabox.width)
                page_height = float(page.mediabox.height)
                
                # Create page number overlay
                number_buffer = io.BytesIO()
                c = canvas.Canvas(number_buffer, pagesize=(page_width, page_height))
                
                c.setFillColor(black)
                c.setFont("Helvetica", 10)
                
                page_text = f"Page {i + 1} of {total_pages}"
                
                # Position based on setting
                y_pos = 30
                if position == "bottom-center":
                    x_pos = page_width / 2
                    c.drawCentredString(x_pos, y_pos, page_text)
                elif position == "bottom-right":
                    x_pos = page_width - 50
                    c.drawRightString(x_pos, y_pos, page_text)
                else:  # bottom-left
                    x_pos = 50
                    c.drawString(x_pos, y_pos, page_text)
                
                c.save()
                number_buffer.seek(0)
                
                # Merge page number with page
                number_reader = PdfReader(number_buffer)
                page.merge_page(number_reader.pages[0])
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            
            logger.info(f"Added page numbers to {total_pages} pages")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error adding page numbers: {e}")
            raise

    @staticmethod
    async def pdf_to_excel(pdf_content: bytes) -> bytes:
        """
        Convert PDF tables to Excel document (XLSX)
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Excel document as bytes
        """
        try:
            import tempfile
            import os
            import pandas as pd
            from openpyxl import Workbook
            
            # Create temporary PDF file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
                pdf_temp.write(pdf_content)
                pdf_path = pdf_temp.name
            
            xlsx_path = pdf_path.replace('.pdf', '.xlsx')
            
            try:
                # Try to extract tables using tabula
                try:
                    import tabula
                    tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
                except Exception:
                    tables = []
                
                # Create Excel workbook
                with pd.ExcelWriter(xlsx_path, engine='openpyxl') as writer:
                    if tables and len(tables) > 0:
                        for i, table in enumerate(tables):
                            if not table.empty:
                                sheet_name = f'Table_{i + 1}'[:31]  # Excel sheet name limit
                                table.to_excel(writer, sheet_name=sheet_name, index=False)
                    else:
                        # If no tables found, create empty sheet with message
                        df = pd.DataFrame({'Note': ['No tables found in PDF. This PDF may contain text/images instead of tabular data.']})
                        df.to_excel(writer, sheet_name='Sheet1', index=False)
                
                # Read the output file
                with open(xlsx_path, 'rb') as f:
                    xlsx_content = f.read()
                
                logger.info("Converted PDF to Excel successfully")
                return xlsx_content
                
            finally:
                # Cleanup temp files
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                if os.path.exists(xlsx_path):
                    os.remove(xlsx_path)
                    
        except Exception as e:
            logger.error(f"Error converting PDF to Excel: {e}")
            raise

    @staticmethod
    async def excel_to_pdf(excel_content: bytes) -> bytes:
        """
        Convert Excel document to PDF
        
        Args:
            excel_content: Excel file content as bytes
            
        Returns:
            PDF file as bytes
        """
        try:
            import pandas as pd
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Read Excel file
            excel_file = io.BytesIO(excel_content)
            xlsx = pd.ExcelFile(excel_file, engine='openpyxl')
            
            # Create PDF
            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
            
            elements = []
            styles = getSampleStyleSheet()
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                
                if df.empty:
                    continue
                
                # Add sheet title
                elements.append(Paragraph(f"<b>{sheet_name}</b>", styles['Heading2']))
                elements.append(Spacer(1, 12))
                
                # Prepare table data
                data = [df.columns.tolist()] + df.values.tolist()
                
                # Convert all values to strings and truncate if too long
                data = [[str(cell)[:50] if cell is not None else '' for cell in row] for row in data]
                
                # Create table
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                
                elements.append(table)
                elements.append(Spacer(1, 24))
            
            doc.build(elements)
            output.seek(0)
            
            logger.info("Converted Excel to PDF successfully")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting Excel to PDF: {e}")
            raise

    @staticmethod
    async def word_to_pdf(word_content: bytes) -> bytes:
        """
        Convert Word document to PDF
        
        Args:
            word_content: Word file content as bytes
            
        Returns:
            PDF file as bytes
        """
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            # Read Word document
            doc = Document(io.BytesIO(word_content))
            
            # Create PDF
            output = io.BytesIO()
            pdf = SimpleDocTemplate(output, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            
            styles = getSampleStyleSheet()
            
            # Create custom styles
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leading=14
            )
            
            elements = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    elements.append(Spacer(1, 6))
                    continue
                
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    elements.append(Paragraph(text, heading_style))
                else:
                    elements.append(Paragraph(text, normal_style))
            
            # Handle tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text[:100] for cell in row.cells]  # Limit cell text
                    table_data.append(row_data)
                
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ]))
                    elements.append(Spacer(1, 12))
                    elements.append(t)
                    elements.append(Spacer(1, 12))
            
            if not elements:
                elements.append(Paragraph("Empty document", normal_style))
            
            pdf.build(elements)
            output.seek(0)
            
            logger.info("Converted Word to PDF successfully")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting Word to PDF: {e}")
            raise
